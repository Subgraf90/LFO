"""
Konverter-Tool: Markdown oder Word → HTML für das LFO-Handbuch

Verwendung:
    # Markdown konvertieren:
    python markdown_to_html_converter.py
    
    # Word-Datei konvertieren:
    python markdown_to_html_converter.py manual_de.docx

Konvertiert manual_de.md oder manual_de.docx zu manual_de.html mit dem 
vordefinierten CSS-Styling. Bilder aus Word werden automatisch extrahiert.

Benötigte Pakete:
    pip install python-docx pillow
"""

import os
import re
import sys
from pathlib import Path
from typing import Optional, Tuple

# Versuche python-docx zu importieren
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Hinweis: python-docx nicht installiert. Word-Unterstützung nicht verfügbar.")
    print("Installieren mit: pip install python-docx pillow")

# Versuche PIL/Pillow zu importieren
try:
    from PIL import Image
    import io
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("Hinweis: Pillow nicht installiert. Bildverarbeitung eingeschränkt.")
    print("Installieren mit: pip install pillow")


def markdown_to_html(markdown_content: str) -> str:
    """
    Konvertiert Markdown-Text zu HTML mit dem LFO-Handbuch-Styling.
    
    Args:
        markdown_content: Markdown-Text als String
        
    Returns:
        Vollständiges HTML-Dokument als String
    """
    
    html = markdown_content
    
    # HTML-Template mit CSS
    html_template = get_html_template()
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>LFO - Handbuch</title>
    <style>
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
        }
        
        h1 {
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-top: 0;
        }
        
        h2 {
            color: #34495e;
            margin-top: 30px;
            border-bottom: 2px solid #ecf0f1;
            padding-bottom: 5px;
        }
        
        h3 {
            color: #555;
            margin-top: 20px;
        }
        
        p {
            margin: 10px 0;
            text-align: justify;
        }
        
        ul, ol {
            margin: 10px 0;
            padding-left: 30px;
        }
        
        li {
            margin: 5px 0;
        }
        
        code {
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Courier New", monospace;
            font-size: 0.9em;
        }
        
        pre {
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            border-left: 4px solid #3498db;
        }
        
        pre code {
            background-color: transparent;
            padding: 0;
        }
        
        blockquote {
            background-color: #e8f4f8;
            border-left: 4px solid #3498db;
            padding: 10px 15px;
            margin: 15px 0;
        }
        
        blockquote.warning {
            background-color: #fff3cd;
            border-left: 4px solid #ffc107;
        }
        
        blockquote.tip {
            background-color: #d4edda;
            border-left: 4px solid #28a745;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        
        th {
            background-color: #3498db;
            color: white;
        }
        
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        
        a {
            color: #3498db;
            text-decoration: none;
        }
        
        a:hover {
            text-decoration: underline;
        }
        
        hr {
            margin-top: 40px;
            border: none;
            border-top: 1px solid #ddd;
        }
        
        img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 20px auto;
            border: 1px solid #ddd;
            border-radius: 5px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        .image-caption {
            text-align: center;
            font-style: italic;
            color: #666;
            margin-top: -15px;
            margin-bottom: 20px;
            font-size: 0.9em;
        }
    </style>
</head>
<body>
{content}
</body>
</html>"""
    
    # Markdown zu HTML Konvertierung
    
    # Überschriften
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2 id="\1">\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    
    # IDs für Überschriften generieren (für Anker-Links)
    def slugify(text):
        text = text.lower()
        text = re.sub(r'[^\w\s-]', '', text)
        text = re.sub(r'[-\s]+', '-', text)
        return text
    
    # IDs zu h2 hinzufügen
    def add_id_to_h2(match):
        title = match.group(1)
        id_slug = slugify(title)
        return f'<h2 id="{id_slug}">{title}</h2>'
    
    html = re.sub(r'<h2>([^<]+)</h2>', add_id_to_h2, html)
    
    # Fettdruck
    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    
    # Code inline
    html = re.sub(r'`([^`]+)`', r'<code>\1</code>', html)
    
    # Code-Blöcke
    html = re.sub(r'```(\w+)?\n(.*?)```', r'<pre><code>\2</code></pre>', html, flags=re.DOTALL)
    
    # Links
    html = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', html)
    
    # Bilder (Markdown-Syntax: ![Alt-Text](pfad.jpg))
    html = re.sub(r'!\[([^\]]*)\]\(([^\)]+)\)', r'<img src="\2" alt="\1">', html)
    
    # Listen
    lines = html.split('\n')
    in_list = False
    list_type = None
    result_lines = []
    
    for line in lines:
        # Ungeordnete Liste
        if re.match(r'^[-*] (.+)$', line):
            if not in_list or list_type != 'ul':
                if in_list:
                    result_lines.append(f'</{list_type}>')
                result_lines.append('<ul>')
                in_list = True
                list_type = 'ul'
            content = re.sub(r'^[-*] (.+)$', r'\1', line)
            result_lines.append(f'<li>{content}</li>')
        # Geordnete Liste
        elif re.match(r'^\d+\. (.+)$', line):
            if not in_list or list_type != 'ol':
                if in_list:
                    result_lines.append(f'</{list_type}>')
                result_lines.append('<ol>')
                in_list = True
                list_type = 'ol'
            content = re.sub(r'^\d+\. (.+)$', r'\1', line)
            result_lines.append(f'<li>{content}</li>')
        # Blockquote (für Hinweise, Tipps, Warnungen)
        elif line.startswith('> '):
            content = line[2:].strip()
            # Spezielle Blockquote-Typen erkennen
            if '**Hinweis:**' in content or '**Wichtig:**' in content:
                result_lines.append(f'<div class="note"><p>{content}</p></div>')
            elif '**Tipp:**' in content:
                result_lines.append(f'<div class="tip"><p>{content}</p></div>')
            elif '**Warnung:**' in content:
                result_lines.append(f'<div class="warning"><p>{content}</p></div>')
            else:
                result_lines.append(f'<blockquote><p>{content}</p></blockquote>')
        # Tabellen
        elif '|' in line and line.strip().startswith('|'):
            if '---' in line:  # Tabellen-Trenner
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if result_lines and '<table>' in result_lines[-1]:
                result_lines.append('<tr>')
                for cell in cells:
                    result_lines.append(f'<td>{cell}</td>')
                result_lines.append('</tr>')
            else:
                result_lines.append('<table>')
                result_lines.append('<tr>')
                for cell in cells:
                    result_lines.append(f'<th>{cell}</th>')
                result_lines.append('</tr>')
        # Horizontale Linie
        elif line.strip() == '---':
            if in_list:
                result_lines.append(f'</{list_type}>')
                in_list = False
            result_lines.append('<hr>')
        # Leerzeile
        elif line.strip() == '':
            if in_list:
                result_lines.append(f'</{list_type}>')
                in_list = False
            result_lines.append('')
        # Normaler Paragraph
        else:
            if in_list:
                result_lines.append(f'</{list_type}>')
                in_list = False
            if line.strip():
                result_lines.append(f'<p>{line}</p>')
    
    if in_list:
        result_lines.append(f'</{list_type}>')
    
    # Tabellen schließen
    html = '\n'.join(result_lines)
    html = re.sub(r'(</tr>)\s*(<tr>)', r'\1\2', html)
    html = re.sub(r'(</tr>)(?!\s*</table>)', r'\1</table>', html)
    
    # Mehrfache Leerzeilen entfernen
    html = re.sub(r'\n{3,}', '\n\n', html)
    
    # Inline-Elemente in Paragraphen verarbeiten
    html = re.sub(r'<p>(.+?)</p>', lambda m: f'<p>{process_inline(m.group(1))}</p>', html)
    
    return html_template.format(content=html)


def process_inline(text: str) -> str:
    """Verarbeitet Inline-Markdown in bereits formatiertem Text."""
    # Fettdruck
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    # Code
    text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
    # Links
    text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', text)
    return text


def word_to_html(docx_path: Path, images_dir: Path) -> Tuple[str, list]:
    """
    Konvertiert eine Word-Datei zu HTML.
    
    Args:
        docx_path: Pfad zur .docx Datei
        images_dir: Verzeichnis für extrahierte Bilder
        
    Returns:
        Tuple von (HTML-Content, Liste der Bild-Pfade)
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx ist nicht installiert. Installieren Sie es mit: pip install python-docx pillow")
    
    doc = Document(str(docx_path))
    html_parts = []
    image_counter = 0
    image_map = {}  # Mapping von Bild-IDs zu Dateinamen
    
    # Bilder-Verzeichnis erstellen
    images_dir.mkdir(exist_ok=True)
    
    # Zuerst alle Bilder extrahieren und mappen
    if hasattr(doc, 'part') and hasattr(doc.part, 'rels'):
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref or (hasattr(rel, 'target_part') and 'image' in str(rel.target_part.content_type)):
                image_counter += 1
                try:
                    # Versuche Bilddaten zu extrahieren
                    if hasattr(rel, 'target_part'):
                        image_blob = rel.target_part.blob
                    else:
                        continue
                    
                    # Dateierweiterung bestimmen
                    ext = 'png'  # Standard
                    if hasattr(rel, 'target_ref'):
                        ext_from_ref = rel.target_ref.split('.')[-1].lower()
                        if ext_from_ref in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                            ext = ext_from_ref
                    elif hasattr(rel, 'target_part') and hasattr(rel.target_part, 'content_type'):
                        ct = rel.target_part.content_type
                        if 'png' in ct:
                            ext = 'png'
                        elif 'jpeg' in ct or 'jpg' in ct:
                            ext = 'jpg'
                        elif 'gif' in ct:
                            ext = 'gif'
                    
                    image_filename = f"image_{image_counter:03d}.{ext}"
                    image_path = images_dir / image_filename
                    
                    # Bild speichern
                    with open(image_path, 'wb') as f:
                        f.write(image_blob)
                    
                    image_map[rel_id] = f"images/{image_filename}"
                except Exception as e:
                    print(f"Warnung: Konnte Bild {rel_id} nicht extrahieren: {e}")
    
    # Jetzt Paragraphen durchgehen und Bilder an der richtigen Stelle einfügen
    for para in doc.paragraphs:
        # Prüfe ob Paragraph Bilder enthält
        has_image = False
        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                # Bild gefunden in diesem Run
                for rel in run._element.xpath('.//a:blip/@r:embed'):
                    if rel in image_map:
                        html_parts.append(f'<img src="{image_map[rel]}" alt="Bild">')
                        has_image = True
        
        # Überschriften erkennen
        if para.style.name.startswith('Heading'):
            level = para.style.name.replace('Heading ', '')
            try:
                level_num = int(level)
                text = para.text.strip()
                if text:
                    if level_num == 1:
                        html_parts.append(f'<h1>{text}</h1>')
                    elif level_num == 2:
                        slug = re.sub(r'[^\w\s-]', '', text.lower())
                        slug = re.sub(r'[-\s]+', '-', slug)
                        html_parts.append(f'<h2 id="{slug}">{text}</h2>')
                    elif level_num == 3:
                        html_parts.append(f'<h3>{text}</h3>')
                    else:
                        html_parts.append(f'<h{level_num}>{text}</h{level_num}>')
            except ValueError:
                # Fallback für normale Paragraphen
                if para.text.strip():
                    html_parts.append(f'<p>{para.text.strip()}</p>')
        else:
            # Normaler Paragraph
            text = para.text.strip()
            if text and not has_image:  # Nur wenn kein Bild im Paragraph ist
                # Formatierung beibehalten (Fettdruck, Kursiv, etc.)
                formatted_text = format_paragraph_runs(para)
                html_parts.append(f'<p>{formatted_text}</p>')
            elif has_image and text:
                # Paragraph mit Bild und Text
                formatted_text = format_paragraph_runs(para)
                html_parts.append(f'<p>{formatted_text}</p>')
    
    image_paths = list(image_map.values())
    
    # Tabellen verarbeiten
    for table in doc.tables:
        html_parts.append('<table>')
        for i, row in enumerate(table.rows):
            html_parts.append('<tr>')
            for cell in row.cells:
                tag = 'th' if i == 0 else 'td'
                cell_text = cell.text.strip()
                html_parts.append(f'<{tag}>{cell_text}</{tag}>')
            html_parts.append('</tr>')
        html_parts.append('</table>')
    
    return '\n'.join(html_parts), image_paths


def format_paragraph_runs(para) -> str:
    """
    Formatiert die Runs eines Paragraphs (Fettdruck, Kursiv, etc.)
    """
    result = []
    for run in para.runs:
        text = run.text
        if run.bold:
            text = f'<strong>{text}</strong>'
        if run.italic:
            text = f'<em>{text}</em>'
        if run.underline:
            text = f'<u>{text}</u>'
        result.append(text)
    return ''.join(result)


def main():
    """Hauptfunktion: Konvertiert manual_de.md oder manual_de.docx zu manual_de.html"""
    script_dir = Path(__file__).parent
    html_file = script_dir / "manual_de.html"
    images_dir = script_dir / "images"
    
    # Prüfe ob Word-Datei als Argument übergeben wurde
    if len(sys.argv) > 1:
        input_file = Path(sys.argv[1])
        if not input_file.is_absolute():
            input_file = script_dir / input_file
    else:
        # Standard: Suche nach .md oder .docx
        md_file = script_dir / "manual_de.md"
        docx_file = script_dir / "manual_de.docx"
        
        if docx_file.exists() and DOCX_AVAILABLE:
            input_file = docx_file
        elif md_file.exists():
            input_file = md_file
        else:
            print(f"Fehler: Weder {md_file} noch {docx_file} gefunden!")
            print("\nVerwendung:")
            print("  python markdown_to_html_converter.py [datei.md oder datei.docx]")
            return
    
    # Bestimme Dateityp
    if input_file.suffix.lower() == '.docx':
        if not DOCX_AVAILABLE:
            print("Fehler: python-docx ist nicht installiert!")
            print("Installieren Sie es mit: pip install python-docx pillow")
            return
        
        print(f"Lese Word-Datei: {input_file}...")
        html_content, image_paths = word_to_html(input_file, images_dir)
        
        if image_paths:
            print(f"✓ {len(image_paths)} Bilder extrahiert nach {images_dir}/")
        
        # HTML-Template hinzufügen
        html_template = get_html_template()
        html_content = html_template.format(content=html_content)
        
    elif input_file.suffix.lower() == '.md':
        print(f"Lese Markdown-Datei: {input_file}...")
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        print("Konvertiere Markdown zu HTML...")
        html_content = markdown_to_html(markdown_content)
    else:
        print(f"Fehler: Unbekannter Dateityp: {input_file.suffix}")
        print("Unterstützte Formate: .md, .docx")
        return
    
    print(f"Speichere {html_file}...")
    with open(html_file, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    print("✓ Konvertierung abgeschlossen!")
    print(f"\nSie können jetzt {html_file} im Browser öffnen oder")
    print("das Help-Fenster im Programm verwenden (F1).")


def get_html_template() -> str:
    """Gibt das HTML-Template zurück."""
    return """<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>LFO - Handbuch</title>
    <style>
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
        }
        
        h1 {
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-top: 0;
        }
        
        h2 {
            color: #34495e;
            margin-top: 30px;
            border-bottom: 2px solid #ecf0f1;
            padding-bottom: 5px;
        }
        
        h3 {
            color: #555;
            margin-top: 20px;
        }
        
        p {
            margin: 10px 0;
            text-align: justify;
        }
        
        ul, ol {
            margin: 10px 0;
            padding-left: 30px;
        }
        
        li {
            margin: 5px 0;
        }
        
        code {
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: "Courier New", monospace;
            font-size: 0.9em;
        }
        
        pre {
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            border-left: 4px solid #3498db;
        }
        
        pre code {
            background-color: transparent;
            padding: 0;
        }
        
        blockquote {
            background-color: #e8f4f8;
            border-left: 4px solid #3498db;
            padding: 10px 15px;
            margin: 15px 0;
        }
        
        blockquote.warning {
            background-color: #fff3cd;
            border-left: 4px solid #ffc107;
        }
        
        blockquote.tip {
            background-color: #d4edda;
            border-left: 4px solid #28a745;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        
        th {
            background-color: #3498db;
            color: white;
        }
        
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        
        a {
            color: #3498db;
            text-decoration: none;
        }
        
        a:hover {
            text-decoration: underline;
        }
        
        hr {
            margin-top: 40px;
            border: none;
            border-top: 1px solid #ddd;
        }
        
        img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 20px auto;
            border: 1px solid #ddd;
            border-radius: 5px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        .image-caption {
            text-align: center;
            font-style: italic;
            color: #666;
            margin-top: -15px;
            margin-bottom: 20px;
            font-size: 0.9em;
        }
    </style>
</head>
<body>
{content}
</body>
</html>"""


if __name__ == "__main__":
    main()

